"""
layout_preserving.py — تصدير DOCX/HTML مع الحفاظ على التخطيط البصري.
يقبل layout_data بصيغة JSON ويُعيد مستنداً يحاكي الكتابة الحاسوبية.

layout_data = {
    "image_path": "path/to/original.jpg",   # اختياري
    "blocks": [
        {"type": "paragraph", "bbox": [x1,y1,x2,y2], "text": "..."},
        {"type": "table",     "bbox": [...], "cells": [[...]]},
        {"type": "image",     "bbox": [...], "image_file": "img.png"},
        {"type": "caption",   "bbox": [...], "text": "..."},
        {"type": "header",    "bbox": [...], "text": "..."},
    ]
}
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _set_rtl(paragraph):
    """ضبط اتجاه الفقرة RTL."""
    pPr = paragraph._element.get_or_add_pPr()
    pPr.set(qn('w:bidi'), '1')


def export_to_docx(layout_data: dict, output_path: str) -> str:
    """
    تصدير layout_data إلى ملف DOCX يحافظ على البنية البصرية.
    يُرجع مسار الملف المُنشأ.
    """
    doc = Document()

    # هوامش موحدة
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # نمط افتراضي RTL
    style = doc.styles['Normal']
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rPr.set(qn('w:rtl'), '1')

    for block in layout_data.get('blocks', []):
        btype = block.get('type', 'paragraph')

        if btype in ('paragraph', 'caption'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)
            run = p.add_run(block.get('text', ''))
            run.font.size = Pt(10 if btype == 'caption' else 12)
            if btype == 'caption':
                run.italic = True

        elif btype == 'header':
            p = doc.add_heading(block.get('text', ''), level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)

        elif btype == 'table':
            cells = block.get('cells', [])
            if not cells:
                continue
            rows, cols = len(cells), max(len(r) for r in cells)
            tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
            tbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for i, row in enumerate(cells):
                for j, cell_text in enumerate(row):
                    if j < cols:
                        c = tbl.cell(i, j)
                        c.text = ''
                        p = c.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        _set_rtl(p)
                        p.add_run(str(cell_text)).font.size = Pt(11)
            doc.add_paragraph()  # مسافة بعد الجدول

        elif btype == 'image':
            img_file = block.get('image_file', '')
            if img_file and os.path.exists(img_file):
                doc.add_picture(img_file, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def ocr_result_to_layout(ocr_json: dict, image_path: str = "") -> dict:
    """
    تحويل مخرجات OCR القياسية إلى تنسيق layout_data.
    """
    layout = {"image_path": image_path, "blocks": []}
    for block in ocr_json.get('blocks', []):
        nb = {
            "type": block.get('type', 'paragraph'),
            "bbox": block.get('bbox', [0, 0, 1, 1]),
            "text": block.get('text', ''),
        }
        if nb["type"] == 'table':
            nb["cells"] = block.get('cells', [])
        elif nb["type"] == 'image':
            nb["image_file"] = block.get('image_file', '')
        layout["blocks"].append(nb)
    return layout
