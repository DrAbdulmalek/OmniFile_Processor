# modules/export/docx_exporter.py
from pathlib import Path
from typing import Optional, List
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from bs4 import BeautifulSoup
import re

class DOCXExporter:
    """مصدّر للنصوص إلى DOCX."""

    def __init__(self, language: str = "ar"):
        self.language = language
        self.is_rtl = language == "ar"

    def export(
        self,
        html_path: Path,
        output_path: Path
    ) -> Path:
        """
        تصدير HTML إلى DOCX.

        Args:
            html_path: مسار ملف HTML.
            output_path: مسار ملف DOCX المخرج.

        Returns:
            Path: مسار ملف DOCX.
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # قراءة HTML
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        # تحليل HTML
        soup = BeautifulSoup(html_content, "html.parser")

        # إنشاء مستند DOCX
        doc = DocxDocument()

        # معالجة الجسم
        body = soup.find("body")
        if body:
            for element in body.children:
                self._process_html_element(element, doc)

        # حفظ الملف
        doc.save(str(output_path))
        return output_path

    def _process_html_element(self, element, doc):
        """معالجة عنصر HTML وإضافته إلى DOCX."""
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "h3":
            doc.add_heading(element.text, level=3)
        elif element.name == "h4":
            doc.add_heading(element.text, level=4)
        elif element.name == "h5":
            doc.add_heading(element.text, level=5)
        elif element.name == "h6":
            doc.add_heading(element.text, level=6)
        elif element.name == "p":
            paragraph = doc.add_paragraph(element.text)
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "table":
            self._process_table(element, doc)
        elif element.name == "img":
            # إضافة صورة (سيتم إضافة دعم لها لاحقًا)
            pass
        elif element.name == "pre":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "strong" or element.name == "b":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.text)
            run.bold = True
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "em" or element.name == "i":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.text)
            run.italic = True
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "s" or element.name == "del":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.text)
            run.font.strike = True
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "br":
            doc.add_paragraph()  # إضافة سطر فارغ
        elif element.name == "hr":
            doc.add_paragraph("---")
        elif element.name == "blockquote":
            paragraph = doc.add_paragraph(element.text)
            paragraph.style = "Block Text"
            if self.is_rtl:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name is None:  # نص عادي
            if element.string and element.string.strip():
                paragraph = doc.add_paragraph(element.string.strip())
                if self.is_rtl:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif element.name == "div" and "metadata" in element.get("class", []):
            # تجاهل قسم البيانات الوصفية
            pass
        else:
            # معالج العناصر غير المعروفة
            if element.string and element.string.strip():
                paragraph = doc.add_paragraph(element.string.strip())
                if self.is_rtl:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def _process_table(self, table_element, doc):
        """معالجة جدول HTML وإضافته إلى DOCX."""
        table = doc.add_table(rows=1, cols=len(table_element.find_all("th", recursive=False)))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # إضافة رأس الجدول
        headers = table_element.find_all("th")
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header.text
            if self.is_rtl:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # إضافة الصفوف
        rows = table_element.find_all("tr")[1:]  # تجاهل رأس الجدول
        for row in rows:
            table_row = table.add_row()
            cells = row.find_all("td")
            for i, cell in enumerate(cells):
                if i < len(table_row.cells):
                    table_row.cells[i].text = cell.text
                    if self.is_rtl:
                        table_row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
